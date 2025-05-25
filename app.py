from flask import Flask, jsonify, render_template, request, redirect, session, flash
import mysql.connector
from sentence_transformers import SentenceTransformer, util
import pandas as pd
import numpy as np
import pickle
import bcrypt
import docx
import PyPDF2
from datetime import datetime
import requests
from flask_sqlalchemy import SQLAlchemy
import re
from io import BytesIO # For PDF processing
import os
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
nltk.download('stopwords')
nltk.download('wordnet')
nltk.download('omw-1.4')
import sqlite3
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.metrics import mean_absolute_error, r2_score
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import StandardScaler
import joblib
import matplotlib.pyplot as plt
import io
import base64

try:
    stop_words = set(stopwords.words('english'))
    lemmatizer = WordNetLemmatizer()
except LookupError:
    print("NLTK data not found. Please run the NLTK download commands.")
    # You might want to exit or raise an error if NLTK data is crucial
    stop_words = set() # Fallback to empty set
    lemmatizer = None # Fallback

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///contacts.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

app.secret_key = 'your_secret_key'
RAPIDAPI_KEY = "236228e410msh8d0f311c7a27b96p1d8798jsnb950f7eb8259"  # Replace with your key
JSEARCH_HEADERS = {
    "X-RapidAPI-Key": RAPIDAPI_KEY,
    "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
}

import sqlite3

# Database Connection Function
def connect_db():
    try:
        conn = sqlite3.connect("final_project.db")  # SQLite creates this file if it doesn't exist
        return conn
    except sqlite3.Error as err:
        print(f"Error: {err}")
        return None

# Use connect_db() Before Executing Queries
conn = connect_db()
if conn:
    cursor = conn.cursor()

    # Create Users Table with Hashed Passwords
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL
        );
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS user_reviews (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            review TEXT NOT NULL
        );
    ''')


    cursor.execute("""
        CREATE TABLE IF NOT EXISTS contact_form (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(255) NOT NULL,
            email VARCHAR(255) NOT NULL,
            message TEXT NOT NULL,
            submitted_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    """)

    conn.commit()
    cursor.close()
    conn.close()
else:
    print("Database connection failed!")

def insert_contact(name, email, message):
    conn = connect_db()
    if conn:
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO contact_submissions (name, email, message)
            VALUES (?, ?, ?);
        ''', (name, email, message))
        conn.commit()
        cursor.close()
        conn.close()
        print("Contact saved successfully!")
    else:
        print("Database connection failed!")






# --- BERT Model Initialization ---
try:
    bert_model = SentenceTransformer("all-MiniLM-L6-v2")
except Exception as e:
    print(f"Error loading SentenceTransformer model: {e}")
    bert_model = None # Handle cases where model loading fails

# --- Predefined Job Descriptions (Fallback) ---
job_descriptions_data = {
    "Data Scientist": "Analyze large volumes of data, build predictive models, and communicate insights to stakeholders. Key skills include Python, R, SQL, machine learning, statistical analysis, and data visualization.",
    "Machine Learning Engineer": "Design, build, and deploy machine learning models and systems for production use. Requires expertise in Python, TensorFlow, PyTorch, scikit-learn, and MLOps practices.",
    "Software Developer": "Develop, test, and maintain software applications, including writing clean and efficient code. Proficiency in languages like Java, C++, Python, or JavaScript, along with software development methodologies.",
    "Web Developer": "Build and maintain websites and web applications, ensuring responsiveness and performance. Skills in HTML, CSS, JavaScript, and frameworks like React, Angular, or Vue.js are essential.",
    "AI Researcher": "Conduct experiments and research in artificial intelligence, publishing findings and advancing the field. Deep knowledge of machine learning algorithms, deep learning, NLP, and research methodologies."
}
core_keywords_for_role = {
    "Data Scientist": {
        "technical_skills": ["python", "r", "sql", "machine learning", "statistical analysis", "data visualization", "deep learning", "nlp", "big data", "spark", "hadoop", "tensorflow", "pytorch", "scikit-learn", "tableau", "power bi"],
        "concepts": ["predictive modeling", "classification", "regression", "clustering", "experimentation", "a/b testing", "data mining", "etl"],
        "soft_skills": ["communication", "problem-solving", "analytical thinking", "storytelling", "collaboration"] # Harder to detect automatically
    },
    "Machine Learning Engineer": {
        "technical_skills": ["python", "java", "c++", "tensorflow", "pytorch", "scikit-learn", "keras", " MLOps", "docker", "kubernetes", "aws", "azure", "gcp", "api development", "data pipelines", "model deployment"],
        "concepts": ["deep learning", "reinforcement learning", "computer vision", "nlp", "model optimization", "scalability", "monitoring"],
        "soft_skills": ["problem-solving", "collaboration", "systems design"]
    },
    "Software Developer": {
        "technical_skills": ["java", "python", "c++", "c#", "javascript", "react", "angular", "vue", "node.js", "spring", ".net", "ruby", "php", "sql", "nosql", "git", "agile", "scrum", "restful apis", "microservices", "data structures", "algorithms"],
        "concepts": ["object-oriented programming", "functional programming", "software development life cycle", "testing", "debugging", "version control"],
        "soft_skills": ["problem-solving", "teamwork", "communication", "attention to detail"]
    },
    "Web Developer": {
        "technical_skills": ["html", "css", "javascript", "typescript", "react", "angular", "vue.js", "jquery", "node.js", "express.js", "php", "laravel", "ruby on rails", "python", "django", "flask", "rest apis", "graphql", "responsive design", "ui/ux principles", "git", "webpack", "babel"],
        "concepts": ["frontend", "backend", "full-stack", "web performance", "seo basics", "accessibility (a11y)"],
        "soft_skills": ["collaboration", "problem-solving", "creativity"]
    },
    "AI Researcher": {
        "technical_skills": ["python", "pytorch", "tensorflow", "machine learning", "deep learning", "nlp", "computer vision", "reinforcement learning", "statistics", "mathematics", "latex", "git"],
        "concepts": ["algorithms", "model architecture", "experimental design", "publication", "research methodology", "state-of-the-art"],
        "soft_skills": ["critical thinking", "innovation", "communication (written & verbal)", "perseverance"]
    },
    "Generic": { # Minimal generic ones if no specific role match
        "technical_skills": [], # Could add "microsoft office", "google suite" if very general
        "concepts": ["project management", "analysis", "reporting"],
        "soft_skills": ["communication", "teamwork", "problem-solving", "leadership", "time management", "adaptability"]
    }
}

# --- Curated YouTube Video Suggestions Database ---
# IDs are placeholders or general examples.
youtube_video_data = {
    "Data Scientist": {
        "low_score": [ # Score < 50
            {"title": "Mistakes to avoid While applying Data Science JObs", "id": "TwqlA5Aj7bo"}, 
            {"title": "Building a Data Science Resume That Gets Interviews", "id": "i5a8ngrAjnI"} 
        ],
        "medium_score": [ # Score 50-75
            {"title": "Tailoring Your Data Science Resume to the Job", "id": "Q_rim0W-IEM"}, # Actual Video
            {"title": "Keywords for Data Scientist Resumes - ATS HACKS", "id": "oQbELXsU9G0"} # Actual Video
        ],
        "general": [ # Always show or if score > 75
            {"title": "Standout Data Science Resume Tips", "id": "zcZRy5SM8i8"} # Actual Video
        ]
    },
    "Software Developer": {
        "low_score": [
            {"title": "Why Your Software Engineer Resume Is Bad (and How to Fix It)", "id": "fShlR-9NqnQ"}, # Actual Video
            {"title": "Resume Mistakes That Keep You From Getting Hired (Software)", "id": "pjqi_M3SPwY"} # Actual Video
        ],
        "medium_score": [
            {"title": "How to Quantify Achievements on a Developer Resume", "id": "YnVTgs3I60g"}, # Actual Video
        ],
        "general": [
            {"title": "The Perfect Software Engineer Resume", "id": "DKG8alY2R34"} 
        ]
    },
    "Machine Learning Engineer": {
        "low_score": [
            {"title": "5 Resume Mistakes to Avoid for ML Jobs", "id": "FHSXA-UBnlM"}, 
        ],
        "medium_score": [
            {"title": "How to Tailor Your Resume for ML Engineer Roles", "id": "Q_rim0W-IEM&t=1s"},  
        ],
        "general": [
            {"title": "How to Build a Machine Learning Engineer Resume", "id": "wlIKCJZEImw"}  # Real video
        ]
    },
    "Web Developer": { # Using generic software dev videos as placeholders for now
        "low_score": [{"title": "Fix Your Web Developer Resume - Top Errors", "id": "bYWWXm-4vQ4"}],
        "medium_score": [{"title": "Highlighting Projects for Web Dev Resumes", "id": "_OWEcsa6-24"}],
        "general": [{"title": "Modern Web Developer Resume Tips", "id": "KZehm-meGMg"}]
    },
    "AI Researcher": { # Using generic data science videos as placeholders
        "low_score": [{"title": "AI Researcher Resume: Common Pitfalls", "id": "FpW8aiJPvts"}],
        "medium_score": [{"title": "Showcasing Research in an AI Resume", "id": "1nUgXnOq-y0&t=175s"}],
        "general": [{"title": "Crafting a Strong AI Researcher CV", "id": "pTN95F3sZG0"}]
    },
    "Generic": { # Fallback for categories not explicitly listed or custom JDs
         "low_score": [
            {"title": "5 Resume Mistakes You NEED to AVOID", "id": "pjqi_M3SPwY&t=35s"}, # Actual Video
            {"title": "How to Write a Resume from Scratch (Beginner's Guide)", "id": "y8YH0Qbu5h4&t=183s"} # Actual Video
        ],
        "medium_score": [
            {"title": "How To Beat Applicant Tracking Systems (ATS)", "id": "hoXuxIt9jK4"}, # Actual Video
            {"title": "Tailor Your Resume to ANY Job in SECONDS", "id": "8DjoABfTz8g"} # Actual Video
        ],
        "general": [
            {"title": "Resume Tips That Will Get You HIRED!", "id": "dQ7Q8ZdnuN0&t=3s"} # Actual Video
        ]
    }
}
# Load Resume Categorization Model
svc_model = pickle.load(open('clf.pkl', 'rb'))
tfidf = pickle.load(open('tfidf.pkl', 'rb'))
le = pickle.load(open('encoder.pkl', 'rb'))

# Helper functions for Resume Text Extraction
def clean_resume(txt):
    cleanText = re.sub('http\S+\s', ' ', txt)
    cleanText = re.sub('RT|cc', ' ', cleanText)
    cleanText = re.sub('#\S+\s', ' ', cleanText)
    cleanText = re.sub('@\S+', '  ', cleanText)
    cleanText = re.sub('[%s]' % re.escape("""!"#$%&'()*+,-./:;<=>?@[\]^_`{|}~"""), ' ', cleanText)
    cleanText = re.sub(r'[^\x00-\x7f]', ' ', cleanText)
    cleanText = re.sub('\s+', ' ', cleanText)
    return cleanText

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ''
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def extract_text_from_txt(file):
    try:
        text = file.read().decode('utf-8')
    except UnicodeDecodeError:
        text = file.read().decode('latin-1')
    return text




# --- Helper Functions ---
def handle_file_upload(uploaded_file):
    # (Keeping this function as it was in the previous "no database" version)
    # ... (same as your last working version) ...
    filename = uploaded_file.filename
    if not filename:
        return "Error: No file selected or filename is empty."
        
    if filename.endswith('.txt'):
        try:
            return uploaded_file.read().decode('utf-8', errors='replace')
        except Exception as e:
            return f"Error reading TXT file: {str(e)}"
    elif filename.endswith('.docx'):
        try:
            import docx
            doc = docx.Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            flash("python-docx library not installed. DOCX processing unavailable.", "warning")
            return "Error: DOCX processing library not available. Please install python-docx."
        except Exception as e:
            return f"Error processing DOCX: {str(e)}"
    elif filename.endswith('.pdf'):
        try:
            from pdfminer.high_level import extract_text # Simpler API for direct extraction
            file_stream = BytesIO(uploaded_file.read())
            return extract_text(file_stream)
        except ImportError:
            flash("pdfminer.six library not installed. PDF processing unavailable.", "warning")
            return "Error: PDF processing library not available. Please install pdfminer.six."
        except Exception as e:
            return f"Error processing PDF: {str(e)}"
    else:
        return "Unsupported file type. Please upload TXT, DOCX, or PDF."
    

def pred(input_resume, confidence_threshold=0.3):
    cleaned_text = clean_resume(input_resume)
    vectorized_text = tfidf.transform([cleaned_text])
    predicted_prob = svc_model.predict_proba(vectorized_text.toarray())[0]
    
    # Filter only high-confidence categories
    high_conf_indices = np.where(predicted_prob >= confidence_threshold)[0]
    
    if len(high_conf_indices) == 0:
        return [("Uncertain", 0.0)]  # if no category passes threshold

    high_conf_categories = le.inverse_transform(high_conf_indices)
    high_conf_scores = predicted_prob[high_conf_indices]
    
    result = list(zip(high_conf_categories, np.round(high_conf_scores * 100, 2)))
    
    # Sort by score descending
    result = sorted(result, key=lambda x: x[1], reverse=True)
    return result

# function for ATS Score 
def calculate_ats_score(resume_text, job_description_text):
    if not bert_model:
        flash("ATS scoring model is not available. Please check server logs.", "error")
        return 0.0
    if not resume_text or not job_description_text:
        return 0.0
    try:
        embeddings = bert_model.encode([resume_text, job_description_text])
        similarity_score = util.cos_sim(embeddings[0], embeddings[1]).item()
        return round(similarity_score * 100, 2)
    except Exception as e:
        print(f"Error in ATS score calculation: {e}")
        flash("Error calculating ATS score. The input might be too large or invalid.", "error")
        return 0.0

def generate_improvement_tips(resume_text, job_description_text, job_category_key_for_core_keywords):
    tips = []
    # Extract keywords from resume and current JD
    resume_keywords_set = extract_keywords_from_text(resume_text)
    job_desc_keywords_set = extract_keywords_from_text(job_description_text)

    # 1. Keywords missing from the SPECIFIC Job Description
    missing_from_jd = sorted(list(job_desc_keywords_set - resume_keywords_set))
    
    # 2. Core keywords for the ROLE missing from the resume
    missing_core_role_keywords = []
    role_specific_core_keywords_data = core_keywords_for_role.get(job_category_key_for_core_keywords, core_keywords_for_role["Generic"])
    
    all_core_keywords_for_role = []
    if role_specific_core_keywords_data: # Check if the key exists
        all_core_keywords_for_role.extend(role_specific_core_keywords_data.get("technical_skills", []))
        all_core_keywords_for_role.extend(role_specific_core_keywords_data.get("concepts", []))
        # Soft skills are harder to match directly this way, can be separate advice

    for core_kw in set(all_core_keywords_for_role): # Use set for efficiency
        if core_kw not in resume_keywords_set:
            missing_core_role_keywords.append(core_kw)
    missing_core_role_keywords = sorted(missing_core_role_keywords)

    # Matched keywords (from JD) - still useful to show
    matched_keywords = sorted(list(job_desc_keywords_set.intersection(resume_keywords_set)))

    # --- Constructing Text Tips (existing logic) ---
    if missing_from_jd: # Prioritize tips about the current JD slightly
        display_missing_count = min(len(missing_from_jd), 5) # Show fewer to make space for core
        if len(missing_from_jd) > 5:
            tips.append(f"From the current job description, consider adding: '{', '.join(missing_from_jd[:display_missing_count])}', and others.")
        else:
            tips.append(f"From the current job description, consider adding: '{', '.join(missing_from_jd[:display_missing_count])}'.")
    
    # ... (your existing logic for action verbs, quantification, team, project) ...
    # This part remains largely the same. For brevity, I'm omitting it but it should be there.
    action_verbs_present = [verb for verb in ['led', 'managed', 'developed', 'created', 'implemented', 'achieved', 'spearheaded', 'optimized', 'innovated', 'designed', 'analyzed'] if verb in resume_text.lower()]
    if len(action_verbs_present) < 3: 
        tips.append("Strengthen your experience descriptions with impactful action verbs to describe your accomplishments (e.g., 'developed', 'managed', 'achieved', 'optimized').")

    if not re.search(r'\d+%', resume_text) and \
       not re.search(r'\$\s?\d+[,.\d]*', resume_text) and \
       not re.search(r'\b\d{3,}\b', resume_text): 
        tips.append("Quantify your achievements wherever possible. Use numbers, percentages, or specific metrics to demonstrate the impact of your work (e.g., 'Increased sales by 15%').")

    if 'team' not in resume_keywords_set and 'collaborat' not in resume_keywords_set: 
        tips.append("If relevant to the role, highlight your teamwork and collaboration skills.")
    if 'project' not in resume_keywords_set and resume_text.lower().count('project') < 2 :
        tips.append("Showcase key projects you've worked on, detailing your specific contributions and the outcomes.")


    # Default tips if others are sparse
    if not tips and matched_keywords:
         tips.append("Your resume shows some alignment! For further refinement, ensure each bullet point clearly states an action, the context/challenge, and a positive result (STAR method).")
    elif not tips and not matched_keywords: # and not missing_from_jd and not missing_core_role_keywords:
        tips.append("Your resume seems significantly different from the job description. Tailor it heavily by focusing on the specific requirements and keywords mentioned in the job posting.")
    elif not tips: 
        tips.append("Review your resume for clarity, conciseness, and impact. Ensure it's tailored to the specific job you're applying for.")

    # Return all three lists: text tips, matched (from JD), missing (from JD), and missing (core for role)
    return tips, matched_keywords, missing_from_jd, missing_core_role_keywords

# Function to extract contact information
def extract_keywords_from_text(text):
    if not text or not lemmatizer: 
        return set()
    text_cleaned = re.sub(r'[^\w\s]', '', text.lower())
    words = text_cleaned.split()
    keywords = {lemmatizer.lemmatize(word) for word in words if word not in stop_words and len(word) > 2}
    return keywords

def extract_contact_info(resume_text):
    if not resume_text:
        return {'Email': 'Not found', 'Phone': 'Not found'}
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b', resume_text)
    phone_match = re.search(r'\b(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{2,4}\)?[-.\s]?)?[\d\s.-]{7,15}\b', resume_text)
    
    email = email_match.group(0) if email_match else 'Not found'
    phone = phone_match.group(0).strip() if phone_match else 'Not found' 

    return {'Email': email, 'Phone': phone}

def get_video_suggestions_for_user(job_category_key, ats_score):
    """
    Selects appropriate video suggestions based on job category and ATS score.
    job_category_key: The key used for analysis (e.g., "Data Scientist", "Generic").
    """
    videos_to_show = []
    
    # Determine the primary category for video lookup
    # If the job_category_key exists in our video_data, use it. Otherwise, use "Generic".
    category_video_set = youtube_video_data.get(job_category_key, youtube_video_data["Generic"])

    score_category = ""
    if ats_score < 50:
        score_category = "low_score"
    elif ats_score < 75:
        score_category = "medium_score"
    
    # Add score-specific videos
    if score_category and score_category in category_video_set:
        videos_to_show.extend(category_video_set[score_category])
        
    # Add general videos from the same category, ensuring no duplicates based on ID
    general_videos = category_video_set.get("general", [])
    existing_ids = {v['id'] for v in videos_to_show}
    for vid in general_videos:
        if vid['id'] not in existing_ids:
            videos_to_show.append(vid)
            existing_ids.add(vid['id'])
            
    # If still no videos (e.g., category had no general and score was high), try Generic general
    if not videos_to_show and job_category_key != "Generic":
        generic_general_videos = youtube_video_data["Generic"].get("general", [])
        for vid in generic_general_videos:
            if vid['id'] not in existing_ids: # Check against already added (should be none here)
                videos_to_show.append(vid)
                existing_ids.add(vid['id'])

    # Limit to a reasonable number, e.g., max 2-3 videos to avoid overwhelming the user
    return videos_to_show[:3]

# Function to save ATS score to SQLite database


# Routes
@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        conn = connect_db()
        if not conn:
            flash("Database connection failed!", "error")
            return redirect('/')

        cursor = conn.cursor()
        cursor.execute("SELECT password_hash FROM users WHERE username = ?", (username,))
        user = cursor.fetchone()
        cursor.close()
        conn.close()

        if user and bcrypt.checkpw(password.encode('utf-8'), user[0].encode('utf-8')):  
            session['user'] = username
            return redirect('/home')
        else:
            flash("Invalid username or password.", "error")

    return render_template('login.html')


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        conn = connect_db()
        if not conn:
            return redirect('/signup')

        cursor = conn.cursor()

        # Check if user already exists
        cursor.execute("SELECT * FROM users WHERE username = ?", (username,))
        existing_user = cursor.fetchone()

        if existing_user:
            flash("Username already exists!", "error")
            cursor.close()
            conn.close()
            return redirect('/signup')

        # Hash the password
        hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())

        # Insert new user (✅ SQLite placeholder fix)
        cursor.execute("INSERT INTO users (username, password_hash) VALUES (?, ?)", 
                       (username, hashed_password.decode('utf-8')))
        conn.commit()

        cursor.close()
        conn.close()

        flash("Signup successful! Please login.", "success")
        return redirect('/')

    return render_template('signup.html')



@app.route('/home')
def home():
    if 'user' not in session:
        return redirect('/login')  # Redirect if the user is not logged in

    # Ensure connection to the database
    conn = connect_db()
    if not conn:
        flash("Failed to connect to the database.", "error")
        return redirect('/login')

    cursor = conn.cursor()

    try:
        # Fetch all reviews from the MySQL database
        cursor.execute("SELECT * FROM user_reviews")
        reviews = cursor.fetchall()
    except mysql.connector.Error as err:
        #flash(f"Database error: {err}", "error")
        reviews = []
    finally:
        cursor.close()  # Always close the cursor after use
        conn.close()  # Always close the connection after use

    # Pass reviews to the template
    return render_template('home.html', username=session['user'], reviews=reviews)



@app.route('/getstart')
def getstart():
    return render_template('getstart.html')

@app.route('/review', methods=['GET', 'POST'])
def review():
    if 'user' not in session:
        return redirect('/login')  # Redirect if not logged in

    if request.method == 'POST':
        user_review = request.form['review']
        username = session['user']

        conn = connect_db()
        try:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO user_reviews (username, review) VALUES (?, ?)", (username, user_review))
            conn.commit()
            flash("Review submitted successfully!", "success")
        except Exception as e:
            conn.rollback()
            flash(f"Database error: {str(e)}", "danger")
        finally:
            conn.close()

        return redirect('/home', username=session['user'])

    return render_template('review.html')



@app.route('/salary-prediction', methods=['GET', 'POST'])
def salary_prediction():
    if 'user' not in session:
        return redirect('/')
    
    if request.method == 'GET':
        # Render form with all dropdown options
        return render_template('salary_prediction.html',
            job_roles=["Data Scientist", "Data Engineer", "Analyst", "Other"],
            industries=["Tech", "Finance", "Healthcare", "Retail", "Other"],
            sectors=["IT", "Finance", "Marketing", "Other"],
            company_sizes=["1-50", "51-200", "201-500", "500+"],
            revenues=["<1M", "1M-10M", "10M-100M", "100M+"],
            locations=["CA", "New York", "Texas", "Florida", "Rome", "Delhi", "Mumbai", "Bangalore"],
            seniorities=["Junior", "Mid", "Senior"]
        )
    
    if request.method == 'POST':
        try:
            # Load model and preprocessor
            model = joblib.load("random_forest_model (2).pkl")
            preprocessor = joblib.load("preprocessor.pkl")
            
            # Process form data
            input_data = {
                "Rating": float(request.form.get('company_rating', 3.5)),
                "Size": request.form.get('company_size', "51-200"),
                "Type of ownership": "Company",
                "Industry": request.form.get('industry', "Tech"),
                "Sector": request.form.get('sector', "IT"),
                "Revenue": request.form.get('revenue', "10M-100M"),
                "Num_comp": int(request.form.get('num_comp', 2)),
                "PerHour": 1 if request.form.get('per_hour') == 'on' else 0,
                "Job Title": request.form.get('job_title', "Data Scientist"),
                "job_state": request.form.get('location', "CA"),
                "Same State": 1 if request.form.get('same_state') == 'on' else 0,
                "Company Age": int(request.form.get('company_age', 5)),
                "Python_yn": 1 if request.form.get('skills_python') == 'on' else 0,
                "Spark": 1 if request.form.get('skills_spark') == 'on' else 0,
                "AWS_yn": 1 if request.form.get('skills_aws') == 'on' else 0,
                "Excel_yn": 1 if request.form.get('skills_excel') == 'on' else 0,
                "desc_len": int(request.form.get('desc_len', 500)),
                "Seniority": request.form.get('job_seniority', "Mid"),
                "Years of Experience": int(request.form.get('experience', 5))
            }
            
            # Create DataFrame and predict
            input_df = pd.DataFrame([input_data])
            input_transformed = preprocessor.transform(input_df)
            prediction = model.predict(input_transformed)[0] * 1000  # Scale up
            
            return render_template('salary_prediction_result.html',
                predicted_salary="${:,.2f}".format(prediction),
                input_data=input_data
            )
            
        except Exception as e:
            return render_template('salary_prediction.html',
                job_roles=["Data Scientist", "Data Engineer", "Analyst", "Other"],
                industries=["Tech", "Finance", "Healthcare", "Retail", "Other"],
                sectors=["IT", "Finance", "Marketing", "Other"],
                company_sizes=["1-50", "51-200", "201-500", "500+"],
                revenues=["<1M", "1M-10M", "10M-100M", "100M+"],
                locations=["CA", "New York", "Texas", "Florida", "Rome", "Delhi", "Mumbai", "Bangalore"],
                seniorities=["Junior", "Mid", "Senior"],
                error=str(e)
            )


@app.route('/resume-categorization', methods=['GET', 'POST'])
def resume_categorization():
    if 'user' not in session:
        return redirect('/')
    
    if request.method == 'POST':
        try:
            uploaded_file = request.files['resume']
            if uploaded_file:
                resume_text = handle_file_upload(uploaded_file)
                top_categories = pred(resume_text)
                return render_template('resume_categorization_result.html', top_categories=top_categories, resume_text=resume_text)
            else:
                flash("Please upload a resume file.", "error")
        except Exception as e:
            flash(f"Error: {str(e)}", "error")
    
    return render_template('resume_categorization.html')


@app.route('/ats-checker', methods=['GET', 'POST'])
def ats_checker_page():
    # Use the correct variable name here
    job_category_keys = list(job_descriptions_data.keys())

    if request.method == 'POST':
        try:
            uploaded_file = request.files.get('resume')
            custom_job_description_text = request.form.get('custom_job_description', '').strip()
            selected_job_category_key = request.form.get('job_category')

            if not uploaded_file or uploaded_file.filename == '':
                flash("Oops! Please upload your resume file.", "error")
                # Pass the correct variable to the template on re-render
                return render_template('ats_checker.html', job_categories=job_category_keys)

            if not custom_job_description_text and not selected_job_category_key:
                flash("Please paste a job description or select a job category to get started.", "error")
                # Pass the correct variable to the template on re-render
                return render_template('ats_checker.html', job_categories=job_category_keys)

            resume_text = handle_file_upload(uploaded_file)
            if resume_text.startswith("Error:") or "Unsupported file type" in resume_text:
                flash(resume_text, "error")
                # Pass the correct variable to the template on re-render
                return render_template('ats_checker.html', job_categories=job_category_keys)

            job_description_to_analyze = ""
            category_for_display_and_videos = "Custom Analysis"
            video_lookup_key = "Generic"

            if custom_job_description_text:
                job_description_to_analyze = custom_job_description_text
                if selected_job_category_key and selected_job_category_key in job_descriptions_data: # Check against new name
                    category_for_display_and_videos = selected_job_category_key
                    video_lookup_key = selected_job_category_key
            elif selected_job_category_key and selected_job_category_key in job_descriptions_data: # Check against new name
                job_description_to_analyze = job_descriptions_data[selected_job_category_key] # Use new name
                category_for_display_and_videos = selected_job_category_key
                video_lookup_key = selected_job_category_key
            else:
                flash("Invalid job input. Please provide a description or select a valid category.", "error")
                # Pass the correct variable to the template on re-render
                return render_template('ats_checker.html', job_categories=job_category_keys)
            
            if not job_description_to_analyze.strip():
                 flash("The job description is empty. Please provide a valid one.", "error")
                 # Pass the correct variable to the template on re-render
                 return render_template('ats_checker.html', job_categories=job_category_keys)

            # ... (rest of the POST logic: ats_score, contact_info, tips, videos) ...
            ats_score = calculate_ats_score(resume_text, job_description_to_analyze)
            contact_info = extract_contact_info(resume_text)
            
            text_only_suggestions, matched_keywords, missing_jd_keywords, missing_core_keywords = [], [], [], []
            
            if ats_score < 98:
                text_only_suggestions, matched_keywords, missing_jd_keywords, missing_core_keywords = \
                    generate_improvement_tips(resume_text, job_description_to_analyze, video_lookup_key)

            relevant_video_suggestions = get_video_suggestions_for_user(video_lookup_key, ats_score)

            return render_template('ats_checker_result.html', 
                                   ats_score=ats_score, 
                                   contact_info=contact_info,
                                   analyzed_category_name=category_for_display_and_videos, 
                                   text_suggestions=text_only_suggestions,
                                   matched_keywords=matched_keywords,
                                   missing_job_description_keywords=missing_jd_keywords,
                                   missing_core_role_keywords=missing_core_keywords,
                                   video_suggestions_list=relevant_video_suggestions)
        
        except Exception as e:
            #print(f"Critical Error in /ats-checker POST: {str(e)}") 
            import traceback
            traceback.print_exc()
            #flash(f"An critical unexpected error occurred. Please check server logs or contact support.", "error")
            # Pass the correct variable to the template on re-render
            return render_template('ats_checker.html', job_categories=job_category_keys)

    # This is for the GET request, make sure job_categories is passed correctly
    return render_template('ats_checker.html', job_categories=job_category_keys)

#  Salary Estimates API
@app.route('/api/salary', methods=['GET'])
def get_salary():
    job_title = request.args.get('job_title', 'Data Scientist')  # Default value
    location = request.args.get('location', 'New York')         # Default value
    
    try:
        url = "https://jsearch.p.rapidapi.com/estimated-salary"
        params = {
            "job_title": job_title,
            "location": location,
            "location_type": "ANY",
            "years_of_experience": "ALL"
        }
        
        response = requests.get(url, headers=JSEARCH_HEADERS, params=params)
        response.raise_for_status()  # Raises HTTPError for bad responses
        
        salary_data = response.json()
        return jsonify({
            "status": "success",
            "data": {
                "job_title": job_title,
                "location": location,
                "min_salary": salary_data.get("min_salary"),
                "max_salary": salary_data.get("max_salary"),
                "median_salary": salary_data.get("median_salary"),
                "currency": salary_data.get("salary_currency") or "USD"
            }
        })
        
    except requests.exceptions.RequestException as e:
        return jsonify({
            "status": "error",
            "message": f"Failed to fetch salary data: {str(e)}"
        }), 500

#  Job Recommendations API       
@app.route('/api/jobs')
def get_jobs():
    # Get parameters with defaults
    query = request.args.get('query', 'Data Science')
    location = request.args.get('location', '')
    page = request.args.get('page', '1')
    
    try:
        # Build API request
        params = {
            "query": f"{query} {location}".strip(),
            "page": page,
            "num_pages": "1",
            "employment_types": "FULLTIME"
        }
        
        headers = {
            "X-RapidAPI-Key": JSEARCH_CONFIG['api_key'],
            "X-RapidAPI-Host": JSEARCH_CONFIG['host']
        }
        
        # Make API call
        response = requests.get(
            JSEARCH_CONFIG['endpoint'],
            headers=headers,
            params=params
        )
        response.raise_for_status()
        
        # Process results
        jobs = response.json().get('data', [])
        processed_jobs = []
        
        for job in jobs:
            processed_jobs.append({
                'title': job.get('job_title'),
                'company': job.get('employer_name'),
                'location': f"{job.get('job_city', '')}, {job.get('job_country', '')}".strip(', '),
                'salary': job.get('job_salary'),
                'posted': format_date(job.get('job_posted_at_datetime_utc')),
                'apply_link': job.get('job_apply_link'),
                'description': job.get('job_description', '').replace('\n', ' ')[:200] + '...'
            })
        
        return jsonify({
            'status': 'success',
            'count': len(processed_jobs),
            'jobs': processed_jobs
        })
        
    except requests.exceptions.HTTPError as e:
        error_detail = e.response.json() if e.response else str(e)
        return jsonify({
            'status': 'error',
            'message': 'JSearch API request failed',
            'detail': error_detail,
            'code': e.response.status_code if e.response else 500
        }), 500
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': 'An unexpected error occurred',
            'detail': str(e)
        }), 500
@app.route('/contact', methods=['GET', 'POST'])
def contact_support():
    if request.method == 'POST':
        # Process form submission
        name = request.form.get('name')
        email = request.form.get('email')
        message = request.form.get('message')
        
        # Here you would typically:
        # 1. Validate inputs
        # 2. Send email/store in database
        # 3. Return success message
        
        return render_template('contact.html', 
                           success=True,
                           name=name)
    
    # GET request - show empty form
    return render_template('contact.html', 
                       success=False)



@app.route('/logout')
def logout():
    session.pop('user', None)
    flash("Logged out successfully!", "success")
    return redirect('/')
@app.route("/submit", methods=["POST"])
def submit_form():
    name = request.form["name"]
    email = request.form["email"]
    message = request.form["message"]
    conn = connect_db()
    cursor = conn.cursor()
    query = "INSERT INTO contact_form (name, email, message) VALUES (?, ?, ?)"
    values = (name, email, message)
    cursor.execute(query, values)
    conn.commit()

    return redirect("/home", username=session['user'])  # or render a success page
if __name__ == '__main__':
    app.run(debug=True)